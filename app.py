import os
import uuid
import base64
from io import BytesIO
import streamlit as st

# ================= PROJECT IMPORTS =================
from speech_to_text import audio_to_text
from text_analyzer import analyze_and_clean_text
from summarizer import structured_summarize
from quiz_generator import generate_quiz, generate_flashcards

api_key = os.getenv("OPENAI_API_KEY")

# ================= PAGE CONFIG =================
st.set_page_config(
    page_title="Lecture Voice-to-Notes Generator",
    layout="wide"
)

# ================= GLOBAL STYLING =================
st.markdown(
    """
    <style>
    .stApp {
        background: linear-gradient(135deg, #e0f2ff 0%, #bae6fd 50%, #7dd3fc 100%);
        color: #0f172a !important;
    }

    .stApp * {
        color: #0f172a !important; 
        font-weight: 500;
    }

    h1, h2, h3, h4, h5, h6 {
        color: #60a5fa !important;
        font-weight: 700;
        text-shadow: 0 0 10px rgba(96, 165, 250, 0.3);
    }

    .quiz-box .question {
        color: #0000ff !important;
        font-weight: 700;
        background: rgba(251, 191, 36, 0.1);
        padding: 8px;
        border-left: 4px solid #fbbf24;
        border-radius: 4px;
        margin: 10px 0;
    }

    .quiz-box .answer {
        color: #4B0082 !important;
        background: rgba(74, 222, 128, 0.1);
        padding: 8px;
        border-left: 4px solid #4ade80;
        border-radius: 4px;
        margin: 5px 0 15px 0;
    }

    div[data-testid="stFileUploader"] {
        background: linear-gradient(135deg, rgba(96, 165, 250, 0.15), rgba(139, 92, 246, 0.15));
        border: 2px solid #60a5fa;
        border-radius: 12px;
        padding: 22px;
    }

    button {
        background: linear-gradient(135deg,#8b5cf6);
        color: white !important;
        border-radius: 8px;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# ================= BACKGROUND IMAGE =================
def add_bg_from_local(image_file):
    if not os.path.exists(image_file):
        return
    with open(image_file, "rb") as img:
        encoded = base64.b64encode(img.read()).decode()

    st.markdown(
        f"""
        <style>
        .stApp {{
            background-image: url("data:image/png;base64,{encoded}");
            background-size: cover;
        }}
        </style>
        """,
        unsafe_allow_html=True
    )

add_bg_from_local("background.png")

# ================= HELPERS =================
def style_quiz_output(raw_text):
    html = '<div class="quiz-box">'
    for line in raw_text.split("\n"):
        if line.startswith("Q"):
            html += f'<div class="question">{line}</div>'
        elif line.startswith("A"):
            html += f'<div class="answer">{line}</div>'
    html += "</div>"
    return html


def download_text(filename, text):
    st.download_button(
        label=f"⬇️ Download {filename}",
        data=text,
        file_name=filename,
        mime="text/plain"
    )

# ================= PDF =================
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

def generate_pdf(summary, quiz, flashcards):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("<b>Lecture Summary</b>", styles["Title"]))
    story.append(Paragraph(summary.replace("\n", "<br/>"), styles["Normal"]))
    story.append(Paragraph("<b>Quiz Questions</b>", styles["Title"]))
    story.append(Paragraph(quiz.replace("\n", "<br/>"), styles["Normal"]))
    story.append(Paragraph("<b>Flashcards</b>", styles["Title"]))
    story.append(Paragraph(flashcards.replace("\n", "<br/>"), styles["Normal"]))

    doc.build(story)
    buffer.seek(0)
    return buffer

# ================= DOCX =================
from docx import Document

def generate_docx(summary, quiz, flashcards):
    doc = Document()
    doc.add_heading("Lecture Summary", level=1)
    doc.add_paragraph(summary)
    doc.add_heading("Quiz Questions", level=1)
    doc.add_paragraph(quiz)
    doc.add_heading("Flashcards", level=1)
    doc.add_paragraph(flashcards)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ================= UI =================
st.title("🎙️ AI Lecture Voice-to-Notes Generator")

uploaded_file = st.file_uploader(
    "Upload Lecture Audio (.wav / .mp3 / .m4a)",
    type=["wav", "mp3", "m4a"]
)

if uploaded_file:
    unique_name = f"{uuid.uuid4()}_{uploaded_file.name}"
    audio_path = os.path.abspath(unique_name)

    with open(audio_path, "wb") as f:
        f.write(uploaded_file.read())

    if st.button("🚀 Generate Notes"):
        with st.spinner("🎤 Transcribing audio..."):
            raw_text = audio_to_text(audio_path)

        with st.spinner("🧹 Cleaning text..."):
            clean_sentences = analyze_and_clean_text(raw_text)

        with st.spinner("📝 Generating summary..."):
            summary = structured_summarize(clean_sentences)
        st.write(summary)

        with st.spinner("❓ Generating quiz..."):
            quiz_output = generate_quiz(summary)
        st.markdown(style_quiz_output(quiz_output), unsafe_allow_html=True)

        with st.spinner("🧠 Generating flashcards..."):
            flashcards = generate_flashcards(summary)
        st.write(flashcards)

        st.download_button("📄 PDF", generate_pdf(summary, quiz_output, flashcards))
        st.download_button("📑 DOCX", generate_docx(summary, quiz_output, flashcards))
